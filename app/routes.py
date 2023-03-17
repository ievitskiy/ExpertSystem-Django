from flask import render_template, flash, redirect, url_for, request, send_from_directory
from flask_login import current_user, login_user, logout_user, login_required
from werkzeug.urls import url_parse
from app import app, db, static_dir
from app.forms import LoginForm, QuestionForm, RegistrationForm
from app.models import User, Post, Quests
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


switch_dict = {'00': '5', '01': '5', '52': '6', '53': '6', '54': '6',
               '62': '7', '63': '7', '64': '7', '70': '8', '71': '8',
               '72': '8', '92': '10', '93': '10', '94': '10'}

switch_end_dict = {'50': '9', '51': '9', '60': '9', '61': '9',
                   '73': '9', '74': '9', '80': '9', '82': '9', '81': '9',
                   '83': '9', '84': '9', '02': '9', '03': '9', '04': '9'}


def db_ref(c_u):
    db.session.add(c_u)
    db.session.commit()


def doc_maker(us_name, fq, sq):
    doc = docx.Document()

    style = doc.styles['Normal']
    style.font.size = Pt(14)
    # добавляем первый параграф
    s1 = doc.add_paragraph()
    s1.add_run(us_name).bold = True
    s1.style = 'Quote'

    # добавляем еще два параграфа
    p = doc.add_paragraph()
    run = p.add_run('Заключение ЭС "Выбор ЯП"')
    run.font.color.rgb = RGBColor(139, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par2 = doc.add_paragraph()
    par2.add_run('\tСледует ' + fq)

    par3 = doc.add_paragraph()
    par3.add_run('Рекомендации по использованию алгоритма:').underline = True

    par4 = doc.add_paragraph()
    par4.add_run('\t' + sq)

    doc.save('../static/report.docx')


@app.shell_context_processor
def make_shell_context():
    return {'db': db, 'User': User, 'Post': Post, 'Quests': Quests}


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('mbase'))


@app.route('/clear')
@login_required
def clear_answers():
    current_user.second_quest_complete = False
    current_user.first_quest_complete = False
    current_user.quest_pos = '0'
    current_user.first_quest = ""
    current_user.second_quest = ""
    db.session.add(current_user)
    db.session.commit()
    return redirect(url_for('question'))


@app.route('/question', endpoint="question", methods=['GET', 'POST'])
@login_required
def question():
    global switch_dict
    global switch_end_dict
    quests = Quests.query.all()
    form = QuestionForm()
    cu = current_user

    # замена значения для упрощения
    for ind, key in switch_dict.items():
        if cu.quest_pos == ind:
            cu.quest_pos = key
            db_ref(cu)
            break

    for i in quests:
        if cu.quest_pos in i.ind.split():
            if not cu.first_quest_complete:
                cu.first_quest = i.quest
                db_ref(cu)
            else:
                cu.second_quest = i.quest
                db_ref(cu)
            break

    if form.validate_on_submit():
        if form.data['submit0']:
            cu.quest_pos += "0"
            db_ref(cu)
            return redirect(url_for('question'))
        if form.data['submit1']:
            cu.quest_pos += "1"
            db_ref(cu)
            print(cu.quest_pos)
            return redirect(url_for('question'))
        if form.data['submit2']:
            cu.quest_pos += "2"
            db_ref(cu)
            return redirect(url_for('question'))
        if form.data['submit3']:
            cu.quest_pos += "3"
            db_ref(cu)
            return redirect(url_for('question'))
        if form.data['subvmit4']:
            cu.quest_pos += "4"
            db_ref(cu)
            return redirect(url_for('question'))

    # замена конечных значений
    for ind, key in switch_end_dict.items():
        if cu.quest_pos == ind:
            cu.quest_pos = key
            db_ref(cu)
        if cu.quest_pos == '9':
            cu.first_quest_complete = True
            for i in quests:
                if cu.quest_pos in i.ind.split():
                    cu.second_quest = i.quest
                    db_ref(cu)
                    break
            break
        if int(cu.quest_pos) in range(100, 105) or int(cu.quest_pos) in [90, 91]:
            cu.second_quest_complete = True
            db_ref(cu)

    return render_template('question.html', form=form)


@app.route('/download', methods=['GET', 'POST'])
@login_required
def download():
    if current_user.first_quest_complete and current_user.second_quest_complete:
        doc_maker(current_user.username, current_user.first_quest, current_user.second_quest)
        return send_from_directory(static_dir, 'report.docx', as_attachment=True)
    else:
        return redirect(url_for('mbase'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    form = RegistrationForm()
    if form.validate_on_submit():
        user = User(username=form.username.data, email=form.email.data)
        user.set_password(form.password.data)
        db.session.add(user)
        db.session.commit()
        flash('Congrats, you are registered!')
        return redirect(url_for('entry'))
    return render_template('register.html', regform=form)


@app.route('/login', endpoint="entry", methods=['GET', 'POST'])
def entry():
    if current_user.is_authenticated:
        return redirect(url_for('mbase'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('Invalid username or password')
            return redirect(url_for('mbase'))
        login_user(user, remember=form.remember_me.data)
        next_page = request.args.get('next')
        if not next_page or url_parse(next_page).netloc != '':
            next_page = url_for('mbase')
        return redirect(next_page)
    return render_template('entry.html', form=form)


@app.route("/", endpoint="mbase")
@app.route('/index', endpoint="mbase")
@app.route('/home', methods=['GET', 'POST'])
def mbase():
    form = LoginForm()
    form2 = RegistrationForm()

    return render_template('homepage.html', form=form, regform=form2)


app.run(debug=True)
